from docx import Document
from io import BytesIO
import streamlit as st
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 금액을 한국어 형식으로 변환하는 함수
def convert_to_korean_currency(amount):
    eok = amount // 100000000  # 억
    man = (amount % 100000000) // 10000  # 만
    result = ""
    if eok > 0:
        result += f"{eok}억 "
    if man > 0 or eok == 0:
        result += f"{man}만원"
    return result.strip()

# docx 파일을 채워주는 함수
def fill_docx_with_results(doc, labor_cost_chief, labor_cost_researcher, labor_cost_assistant_researcher, labor_cost_assistant, total_labor_cost, 
                           travel_expense, printing_expense, meeting_expense, computing_expense, total_expenses, profit_rate,
                           overhead, profit, total_cost, vat, calculated_total_project_cost, labor_cost_ratio, expenses_ratio, overhead_rate):

    # 값을 입력해야 하는 구역을 찾아서 업데이트
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text

                # 인건비 관련 업데이트
                if 'a1' in cell_text:
                    cell.text = cell_text.replace('a1', f"{int(labor_cost_chief):,}")
                elif 'a2' in cell_text:
                    cell.text = cell_text.replace('a2', f"{int(labor_cost_researcher):,}")
                elif 'a3' in cell_text:
                    cell.text = cell_text.replace('a3', f"{int(labor_cost_assistant_researcher):,}")
                elif 'a0' in cell_text:
                    cell.text = cell_text.replace('a0', f"{int(total_labor_cost):,}")

                # 경비 관련 업데이트
                if 'b2' in cell_text:
                    cell.text = cell_text.replace('b2', f"{int(travel_expense):,}")
                elif 'b3' in cell_text:
                    cell.text = cell_text.replace('b3', f"{int(printing_expense):,}")
                elif 'b4' in cell_text:
                    cell.text = cell_text.replace('b4', f"{int(meeting_expense):,}")
                elif 'b5' in cell_text:
                    cell.text = cell_text.replace('b5', f"{int(computing_expense):,}")
                elif 'b1' in cell_text:
                    cell.text = cell_text.replace('b1', f"{int(total_expenses):,}")

                # 원가와 기타 항목 업데이트
                if 'c1' in cell_text:
                    cell.text = cell_text.replace('c1', f"{int(overhead):,}")
                elif 'c2' in cell_text:
                    cell.text = cell_text.replace('c2', f"{int(profit):,}")
                elif 'c3' in cell_text:
                    cell.text = cell_text.replace('c3', f"{int(total_cost):,}")
                elif 'c4' in cell_text:
                    cell.text = cell_text.replace('c4', f"{int(vat):,}")
                elif 'c5' in cell_text:
                    cell.text = cell_text.replace('c5', f"{int(calculated_total_project_cost):,}")

                # 비율 업데이트
                if 'd1' in cell_text:
                    cell.text = cell_text.replace('d1', f"{labor_cost_ratio*100:.2f}%")
                elif 'd2' in cell_text:
                    cell.text = cell_text.replace('d2', f"{expenses_ratio*100:.2f}%")
                elif 'd3' in cell_text:
                    cell.text = cell_text.replace('d3', f"{overhead_rate:.2f}%")
                elif 'd4' in cell_text:
                    cell.text = cell_text.replace('d4', f"{profit_rate:.2f}%")

                # 글자 크기 및 가운데 정렬 설정
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 가운데 정렬
                    for run in paragraph.runs:
                        run.font.size = Pt(12)  # 글자 크기 12pt로 설정

    # Streamlit에서 파일로 저장
    output = BytesIO()
    doc.save(output)
    return output

# Streamlit 애플리케이션 메인 함수
def main():
    st.title("정책연구용역 원가계산 프로그램")

    # 세션 상태 초기화
    if 'output' not in st.session_state:
        st.session_state['output'] = None
    if 'calculation_done' not in st.session_state:
        st.session_state['calculation_done'] = False
    if 'doc' not in st.session_state:
        st.session_state['doc'] = None

    # 서식 파일 업로드
    uploaded_template = st.file_uploader("서식 파일 업로드 (.docx)", type="docx")
    
    if uploaded_template:
        # 업로드된 docx 객체를 세션 상태에 저장
        st.session_state['doc'] = Document(uploaded_template)

        # 사용자 입력
        total_project_cost = st.number_input("총사업비 (원)", min_value=0, step=1000000, format="%d")
        
        # 총사업비를 한국어 금액으로 표시
        if total_project_cost > 0:
            korean_total_project_cost = convert_to_korean_currency(total_project_cost)
            st.write(f"입력한 총사업비: {korean_total_project_cost}")

        # 연구기간 입력
        research_period = st.number_input("연구기간 (개월)", min_value=1, step=1)

        # 추천 연구기간 표시
        if total_project_cost > 0:
            if total_project_cost <= 50000000:
                recommended_period = "5000만원 이하 추천 연구기간: 3~6개월"
            elif total_project_cost <= 100000000:
                recommended_period = "1억원 이하 추천 연구기간: 6~9개월"
            else:
                recommended_period = "1억원 이상 추천 연구기간: 9~12개월"
            st.markdown(f"<span style='color:blue; font-weight:bold;'>{recommended_period}</span>", unsafe_allow_html=True)

        st.header("연구인원 입력")
        col1, col2 = st.columns(2)
        with col1:
            num_chief = st.number_input("책임연구원 수", min_value=1, max_value=1, value=1, step=0)
            num_assistant_researcher = st.number_input("연구보조원 수", min_value=0, step=1)
        with col2:
            num_researcher = st.number_input("연구원 수", min_value=0, step=1)
            num_assistant = st.number_input("보조원 수", min_value=0, step=1)


        # 추천 연구인원 표시
        if total_project_cost > 0:
            if total_project_cost <= 50000000:
                recommended_staff = """
                **추천 연구인원**:
                - 책임연구원: 1명
                - 연구원: 1~2명
                - 연구보조원: 1~2명
                """
            elif total_project_cost <= 100000000:
                recommended_staff = """
                **추천 연구인원**:
                - 책임연구원: 1명
                - 연구원: 2~3명
                - 연구보조원: 2~3명
                """
            else:
                recommended_staff = """
                **추천 연구인원**:
                - 책임연구원: 1명
                - 연구원: 3~5명
                - 연구보조원: 3~5명
                """
            st.markdown(f"<div style='color:blue; font-weight:bold;'>{recommended_staff}</div>", unsafe_allow_html=True)

        st.header("일반관리비 및 이윤 비율 입력")
        col1, col2 = st.columns(2)
        with col1:
            overhead_rate = st.slider("일반관리비 비율 (%) (최대 6%)", min_value=0.00, max_value=6.00, value=5.00, step=0.1)
        with col2:
            profit_rate = st.slider("이윤 비율 (%) (최대 10%)", min_value=0.00, max_value=10.00, value=5.00, step=0.1)

        # 경비 입력
        st.header("경비 입력")
        with st.expander("경비 입력", expanded=True):
            tab1, tab2, tab3, tab4 = st.tabs(["출장비", "인쇄비", "전산처리비", "회의비"])

            # 1. 출장비 입력
            with tab1:
                st.subheader("출장비 입력")
                col1, col2 = st.columns(2)
                with col1:
                    num_trips_per_month = st.number_input("월별 출장 횟수", min_value=0, step=1, value=6)
                    transportation_cost_per_trip = st.number_input("운임비 (1회당, 원, 최대 100,000원)", min_value=0, max_value=100000, step=1000, value=20000)
                total_personnel = num_chief + num_researcher + num_assistant_researcher + num_assistant
                total_trips_per_person = num_trips_per_month * research_period
                total_trips = total_trips_per_person * total_personnel
                daily_allowance = 25000  # 일비
                meal_allowance = 25000  # 식비
                travel_expense = total_trips * (daily_allowance + meal_allowance + transportation_cost_per_trip)

            # 2. 인쇄비 입력
            with tab2:
                st.subheader("인쇄비 입력")
                printing_expense = 0
                col1, col2 = st.columns(2)
                with col1:
                    include_inception_meeting = st.checkbox("착수보고회 포함", value=True)
                    include_interim_meeting = st.checkbox("중간보고회 포함", value=True)
                    include_final_meeting = st.checkbox("최종보고회 포함", value=True)
                with col2:
                    fixed_unit_price = 10000  # 단가 1만원 고정
                    if include_inception_meeting:
                        inception_copies = st.number_input("착수보고회 부수", min_value=0, step=1, value=30)
                        inception_cost = fixed_unit_price * inception_copies
                        printing_expense += inception_cost
                    if include_interim_meeting:
                        interim_copies = st.number_input("중간보고회 부수", min_value=0, step=1, value=30)
                        interim_cost = fixed_unit_price * interim_copies
                        printing_expense += interim_cost
                    if include_final_meeting:
                        final_copies = st.number_input("최종보고회 부수", min_value=0, step=1, value=30)
                        final_cost = fixed_unit_price * final_copies
                        printing_expense += final_cost

            # 3. 전산처리비 입력
            with tab3:
                st.subheader("전산처리비 입력")
                computing_expense = 0
                col1, col2 = st.columns(2)
                with col1:
                    include_copy_paper = st.checkbox("복사용지 포함", value=True)
                with col2:
                    include_toner = st.checkbox("토너 포함", value=True)
                if include_copy_paper:
                    col1, col2 = st.columns(2)
                    with col1:
                        copy_paper_unit_price = st.number_input("복사용지 단가 (원)", min_value=0, step=1000, value=20000)
                    with col2:
                        copy_paper_quantity = st.number_input("복사용지 수량 (박스)", min_value=0, step=1, value=1)
                    computing_expense += copy_paper_unit_price * copy_paper_quantity
                if include_toner:
                    col1, col2 = st.columns(2)
                    with col1:
                        toner_unit_price = st.number_input("토너 단가 (원)", min_value=0, step=1000, value=200000)
                    with col2:
                        toner_quantity = st.number_input("토너 수량 (개)", min_value=0, step=1, value=1)
                    computing_expense += toner_unit_price * toner_quantity

            # 4. 회의비 입력
            with tab4:
                st.subheader("회의비 입력")
                meeting_expense = 0
                meeting_unit_price = 120000  # 고정 단가
                col1, col2 = st.columns(2)
                with col1:
                    include_advisory_meeting = st.checkbox("자문회의 포함", value=True)
                    include_forum = st.checkbox("토론회 포함", value=True)
                    include_public_hearing = st.checkbox("공청회 포함", value=True)
                with col2:
                    if include_advisory_meeting:
                        advisory_attendees = st.number_input("자문회의 참석인원", min_value=0, step=1, value=5)
                        meeting_expense += meeting_unit_price * advisory_attendees
                    if include_forum:
                        forum_attendees = st.number_input("토론회 참석인원", min_value=0, step=1, value=5)
                        meeting_expense += meeting_unit_price * forum_attendees
                    if include_public_hearing:
                        public_hearing_attendees = st.number_input("공청회 참석인원", min_value=0, step=1, value=5)
                        meeting_expense += meeting_unit_price * public_hearing_attendees

        # 총 경비 계산
        total_expenses = travel_expense + printing_expense + computing_expense + meeting_expense

        if st.button("계산하기"):
            max_participation_chief = 50.0
            max_participation_researcher = 80.0
            max_participation_assistant_researcher = 80.0
            max_participation_assistant = 100.0

            salary_chief = 8000000
            salary_researcher = 6000000
            salary_assistant_researcher = 4000000
            salary_assistant = 3000000

            low = 0.0
            high = 1.0  # 조정 비율 (0 ~ 1)
            tolerance = 999  # 천원 단위 이내
            max_iterations = 1000

            success = False

            for _ in range(max_iterations):
                adjustment_factor = (low + high) / 2

                participation_chief = max_participation_chief * adjustment_factor
                participation_researcher = max_participation_researcher * adjustment_factor
                participation_assistant_researcher = max_participation_assistant_researcher * adjustment_factor
                participation_assistant = max_participation_assistant * adjustment_factor

                labor_cost_chief = num_chief * salary_chief * (participation_chief / 100) * research_period
                labor_cost_researcher = num_researcher * salary_researcher * (participation_researcher / 100) * research_period
                labor_cost_assistant_researcher = num_assistant_researcher * salary_assistant_researcher * (participation_assistant_researcher / 100) * research_period
                labor_cost_assistant = num_assistant * salary_assistant * (participation_assistant / 100) * research_period
                total_labor_cost = labor_cost_chief + labor_cost_researcher + labor_cost_assistant_researcher + labor_cost_assistant

                direct_cost = total_labor_cost + total_expenses
                overhead = direct_cost * (overhead_rate / 100)
                profit = (direct_cost + overhead) * (profit_rate / 100)
                total_cost = direct_cost + overhead + profit
                vat = total_cost * 0.10
                calculated_total_project_cost = total_cost + vat

                difference = calculated_total_project_cost - total_project_cost

                if 0 <= difference <= tolerance:
                    success = True
                    break

                if difference < 0:
                    low = adjustment_factor
                else:
                    high = adjustment_factor

            if not success:
                st.error("조건을 만족하는 참여율을 찾을 수 없습니다.")
                st.stop()

            labor_cost_ratio = total_labor_cost / total_cost
            expenses_ratio = total_expenses / total_cost

            error_flag = False
            if labor_cost_ratio > 0.70:
                st.error(f"인건비가 총원가의 70%를 초과합니다. 현재 비율: {labor_cost_ratio*100:.2f}%")
                error_flag = True

            if expenses_ratio < 0.14:
                st.error(f"경비가 총원가의 14% 미만입니다. 현재 비율: {expenses_ratio*100:.2f}%")
                error_flag = True

            if error_flag:
                st.stop()

            # 계산 결과를 세션 상태에 저장
            st.session_state['calculation_results'] = {
                'labor_cost_chief': labor_cost_chief,
                'labor_cost_researcher': labor_cost_researcher,
                'labor_cost_assistant_researcher': labor_cost_assistant_researcher,
                'labor_cost_assistant': labor_cost_assistant,
                'total_labor_cost': total_labor_cost,
                'travel_expense': travel_expense,
                'printing_expense': printing_expense,
                'computing_expense': computing_expense,
                'meeting_expense': meeting_expense,
                'total_expenses': total_expenses,
                'overhead': overhead,
                'profit': profit,
                'total_cost': total_cost,
                'vat': vat,
                'calculated_total_project_cost': calculated_total_project_cost,
                'participation_chief': participation_chief,
                'participation_researcher': participation_researcher,
                'participation_assistant_researcher': participation_assistant_researcher,
                'participation_assistant': participation_assistant,
                'labor_cost_ratio': labor_cost_ratio,
                'expenses_ratio': expenses_ratio,
                'direct_cost': direct_cost,
                'overhead_rate': overhead_rate,
                'profit_rate': profit_rate
            }

            st.session_state['calculation_done'] = True

        if st.session_state['calculation_done']:
            # 세션 상태에서 결과 가져오기
            results = st.session_state['calculation_results']
            labor_cost_chief = results['labor_cost_chief']
            labor_cost_researcher = results['labor_cost_researcher']
            labor_cost_assistant_researcher = results['labor_cost_assistant_researcher']
            labor_cost_assistant = results['labor_cost_assistant']
            total_labor_cost = results['total_labor_cost']
            travel_expense = results['travel_expense']
            printing_expense = results['printing_expense']
            computing_expense = results['computing_expense']
            meeting_expense = results['meeting_expense']
            total_expenses = results['total_expenses']
            overhead = results['overhead']
            profit = results['profit']
            total_cost = results['total_cost']
            vat = results['vat']
            calculated_total_project_cost = results['calculated_total_project_cost']
            participation_chief = results['participation_chief']
            participation_researcher = results['participation_researcher']
            participation_assistant_researcher = results['participation_assistant_researcher']
            participation_assistant = results['participation_assistant']
            labor_cost_ratio = results['labor_cost_ratio']
            expenses_ratio = results['expenses_ratio']
            direct_cost = results['direct_cost']
            overhead_rate = results['overhead_rate']
            profit_rate = results['profit_rate']

            # 결과 출력
            st.header("계산 결과")
            st.subheader("인건비 상세내역")
            st.write(f"책임연구원 인건비: {int(labor_cost_chief)}원 (참여율: {participation_chief:.2f}%)")
            st.write(f"연구원 인건비: {int(labor_cost_researcher)}원 (참여율: {participation_researcher:.2f}%)")
            st.write(f"연구보조원 인건비: {int(labor_cost_assistant_researcher)}원 (참여율: {participation_assistant_researcher:.2f}%)")
            st.write(f"보조원 인건비: {int(labor_cost_assistant)}원 (참여율: {participation_assistant:.2f}%)")
            st.write(f"총 인건비: {int(total_labor_cost)}원")

            st.subheader("경비 상세내역")
            st.write(f"출장비: {int(travel_expense)}원")
            st.write(f"인쇄비: {int(printing_expense)}원")
            st.write(f"전산처리비: {int(computing_expense)}원")
            st.write(f"회의비: {int(meeting_expense)}원")
            st.write(f"총 경비: {int(total_expenses)}원")

            st.subheader("원가 계산")
            st.write(f"순원가 (인건비 + 경비): {int(direct_cost)}원")
            st.write(f"일반관리비 ({overhead_rate:.1f}%): {int(overhead)}원")
            st.write(f"이윤 ({profit_rate:.1f}%): {int(profit)}원")
            st.write(f"총원가: {int(total_cost)}원")
            st.write(f"부가가치세 (10%): {int(vat)}원")
            st.write(f"총사업비: {int(calculated_total_project_cost)}원")

            st.subheader("비율 검증")
            st.write(f"인건비 비율: {labor_cost_ratio*100:.2f}% (70% 이하)")
            st.write(f"경비 비율: {expenses_ratio*100:.2f}% (14% 이상)")

            # DOCX 파일 생성 버튼
            if st.button("결과를 DOCX 파일로 저장하기"):
                output = fill_docx_with_results(
                    doc=st.session_state['doc'],
                    labor_cost_chief=labor_cost_chief,
                    labor_cost_researcher=labor_cost_researcher,
                    labor_cost_assistant_researcher=labor_cost_assistant_researcher,
                    labor_cost_assistant=labor_cost_assistant,
                    total_labor_cost=total_labor_cost,
                    travel_expense=travel_expense,
                    printing_expense=printing_expense,
                    meeting_expense=meeting_expense,
                    computing_expense=computing_expense,
                    total_expenses=total_expenses,
                    profit_rate=profit_rate,
                    overhead=overhead,
                    profit=profit,
                    total_cost=total_cost,
                    vat=vat,
                    calculated_total_project_cost=calculated_total_project_cost,
                    labor_cost_ratio=labor_cost_ratio,
                    expenses_ratio=expenses_ratio,
                    overhead_rate=overhead_rate
                )
                st.session_state['output'] = output.getvalue()
                st.success("파일이 생성되었습니다. 아래의 버튼을 클릭하여 다운로드하세요.")

        # 다운로드 버튼 표시
        if st.session_state['output'] is not None:
            st.download_button(
                label="DOCX 다운로드",
                data=st.session_state['output'],
                file_name="계산결과.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            # 다시 시작하기 버튼
            if st.button("다시 시작하기"):
                st.session_state['output'] = None
                st.session_state['calculation_done'] = False
                st.session_state['doc'] = None
                st.experimental_rerun()

if __name__ == "__main__":
    main()
